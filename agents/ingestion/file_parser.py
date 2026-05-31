"""
Parses raw input files (Excel, PDF, Word, CSV, Markdown, text, images) into plain text.
Images are described using GPT-4o vision. All other formats use local libraries.
"""
from __future__ import annotations

import base64
import csv
import io
from pathlib import Path

import openai

SUPPORTED = {
    ".xlsx", ".xls",           # Excel
    ".pdf",                    # PDF
    ".docx", ".doc",           # Word
    ".csv",                    # CSV
    ".md", ".txt", ".rst",     # Plain text
    ".png", ".jpg", ".jpeg", ".webp",  # Images (GPT-4o vision)
}

IMAGE_TYPES = {".png", ".jpg", ".jpeg", ".webp"}
MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}


def parse_file(path: Path, client: openai.OpenAI) -> tuple[str, dict]:
    """
    Returns (content_text, metadata_dict).
    content_text is the full extracted/described text.
    metadata contains file-specific info (sheet names, page count, etc.).
    """
    ext = path.suffix.lower()

    if ext in (".xlsx", ".xls"):
        return _parse_excel(path)
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext in (".docx", ".doc"):
        return _parse_docx(path)
    if ext == ".csv":
        return _parse_csv(path)
    if ext in (".md", ".txt", ".rst"):
        return _parse_text(path)
    if ext in IMAGE_TYPES:
        return _parse_image(path, client)

    return f"[Unsupported file type: {ext}]", {}


def _parse_excel(path: Path) -> tuple[str, dict]:
    try:
        import openpyxl
    except ImportError:
        return "[openpyxl not installed — cannot parse Excel]", {}

    wb = openpyxl.load_workbook(path, data_only=True)
    lines: list[str] = []
    sheet_names = wb.sheetnames

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        lines.append(f"\n=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                lines.append(" | ".join(cells))

    return "\n".join(lines), {"sheets": sheet_names, "sheet_count": len(sheet_names)}


def _parse_pdf(path: Path) -> tuple[str, dict]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[pypdf not installed — cannot parse PDF]", {}

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text}")

    return "\n\n".join(pages), {"page_count": len(reader.pages)}


def _parse_docx(path: Path) -> tuple[str, dict]:
    try:
        import docx as python_docx
    except ImportError:
        return "[python-docx not installed — cannot parse Word]", {}

    doc = python_docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables_text: list[str] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        tables_text.append("\n".join(rows))

    content = "\n".join(paragraphs)
    if tables_text:
        content += "\n\n[Tables]\n" + "\n\n".join(tables_text)

    return content, {"paragraph_count": len(paragraphs), "table_count": len(doc.tables)}


def _parse_csv(path: Path) -> tuple[str, dict]:
    rows: list[str] = []
    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows), {"row_count": len(rows)}


def _parse_text(path: Path) -> tuple[str, dict]:
    content = path.read_text(encoding="utf-8", errors="replace")
    return content, {"char_count": len(content)}


def _parse_image(path: Path, client: openai.OpenAI) -> tuple[str, dict]:
    ext = path.suffix.lower()
    media_type = MEDIA_TYPES.get(ext, "image/png")

    with open(path, "rb") as f:
        image_b64 = base64.standard_b64encode(f.read()).decode("utf-8")

    response = client.chat.completions.create(
        model="gpt-4o",
        max_tokens=1500,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{media_type};base64,{image_b64}"},
                },
                {
                    "type": "text",
                    "text": (
                        "You are extracting content from a research input image. "
                        "Describe everything visible: charts (with axis labels and values), "
                        "tables (extract all rows/columns), diagrams (describe structure), "
                        "handwritten notes (transcribe), screenshots (extract key data). "
                        "Be exhaustive — this is used for research synthesis."
                    ),
                },
            ],
        }],
    )
    description = response.choices[0].message.content or "[Could not describe image]"
    return description, {"media_type": media_type, "size_bytes": path.stat().st_size}
